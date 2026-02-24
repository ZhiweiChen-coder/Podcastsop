from __future__ import annotations

from io import BytesIO
from typing import Optional


def export_docx_bytes(*, text: str, title: Optional[str] = None) -> bytes:
    """
    Create a simple .docx from plain text.
    Requires dependency: python-docx
    """
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise RuntimeError("Missing dependency for DOCX export. Please install: python-docx") from e

    doc = Document()
    if title and title.strip():
        doc.add_heading(title.strip(), level=1)

    # Preserve paragraphs by blank-line splitting; keep line breaks within paragraphs.
    blocks = (text or "").replace("\r\n", "\n").split("\n\n")
    for block in blocks:
        block = block.strip("\n")
        if not block:
            continue
        p = doc.add_paragraph()
        lines = block.split("\n")
        for idx, line in enumerate(lines):
            if idx > 0:
                p.add_run().add_break()
            p.add_run(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

